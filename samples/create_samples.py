"""
create_samples.py
디스플레이 공장 정비/설비 관리 샘플 파일 15개 생성
  - Excel 5개 (openpyxl)
  - Word  5개 (python-docx)
  - PDF   5개 (fpdf2)
"""

import os
from pathlib import Path

SAVE_DIR = Path(__file__).parent
SAVE_DIR.mkdir(parents=True, exist_ok=True)

# ──────────────────────────────────────────────
# 1. Excel 샘플 5개
# ──────────────────────────────────────────────
def create_excel_samples():
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    thin = Side(style="thin")
    border_all = Border(left=thin, right=thin, top=thin, bottom=thin)

    def header_fill(color="1F4E79"):
        return PatternFill("solid", fgColor=color)

    def set_col_widths(ws, widths):
        for i, w in enumerate(widths, 1):
            ws.column_dimensions[get_column_letter(i)].width = w

    def apply_border(ws, min_row, max_row, min_col, max_col):
        for row in ws.iter_rows(min_row=min_row, max_row=max_row,
                                min_col=min_col, max_col=max_col):
            for cell in row:
                cell.border = border_all

    # ── 1) 정비비용정산서 ──────────────────────
    wb = Workbook()
    ws = wb.active
    ws.title = "정비비용정산서"

    ws["A1"] = "정비비용정산서"
    ws["A1"].font = Font(size=16, bold=True, color="FFFFFF")
    ws["A1"].fill = header_fill()
    ws["A1"].alignment = Alignment(horizontal="center")
    ws.merge_cells("A1:G1")

    ws["A2"] = "작성일: 2026-05-09"
    ws["D2"] = "담당: 김정비"
    ws["G2"] = "결재: 이과장"

    headers = ["일자", "설비코드", "품명", "단가(원)", "수량", "합계(원)", "비고"]
    for c, h in enumerate(headers, 1):
        cell = ws.cell(row=4, column=c, value=h)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = header_fill("2E75B6")
        cell.alignment = Alignment(horizontal="center")
        cell.border = border_all

    data = [
        ("2026-05-02", "EQ-A101", "볼베어링 6204", 8500,  4, None, ""),
        ("2026-05-03", "EQ-B205", "오링 세트",     3200, 10, None, ""),
        ("2026-05-04", "EQ-A101", "V벨트 B-50",   12000,  2, None, ""),
        ("2026-05-05", "EQ-C310", "필터엘리먼트",  25000,  3, None, ""),
        ("2026-05-06", "EQ-B205", "실리콘그리스",   4500,  5, None, ""),
        ("2026-05-07", "EQ-D402", "냉각팬모터",   185000,  1, None, "외주"),
        ("2026-05-08", "EQ-A101", "유압호스 1m",   9800,  6, None, ""),
        ("2026-05-09", "EQ-C310", "솔레노이드밸브",98000,  2, None, ""),
    ]

    subtotals = {}  # 설비코드별 소계
    for r, (date, eq, name, unit, qty, _, remark) in enumerate(data, 5):
        total = unit * qty
        ws.cell(row=r, column=1, value=date)
        ws.cell(row=r, column=2, value=eq)
        ws.cell(row=r, column=3, value=name)
        ws.cell(row=r, column=4, value=unit)
        ws.cell(row=r, column=5, value=qty)
        ws.cell(row=r, column=6, value=total)
        ws.cell(row=r, column=7, value=remark)
        subtotals[eq] = subtotals.get(eq, 0) + total

    # 소계행
    sub_row = 5 + len(data)
    ws.cell(row=sub_row, column=3, value="소계").font = Font(bold=True)
    grand = sum(subtotals.values())
    ws.cell(row=sub_row, column=6, value=grand).font = Font(bold=True)
    ws.cell(row=sub_row, column=6).fill = PatternFill("solid", fgColor="FFF2CC")

    apply_border(ws, 4, sub_row, 1, 7)
    set_col_widths(ws, [12, 12, 18, 12, 8, 14, 10])
    wb.save(SAVE_DIR / "정비비용정산서.xlsx")
    print("  [Excel 1/5] 정비비용정산서.xlsx 생성 완료")

    # ── 2) 설비점검일지 ──────────────────────
    wb = Workbook()
    ws = wb.active
    ws.title = "설비점검일지"

    ws["A1"] = "설비 일상점검일지"
    ws["A1"].font = Font(size=15, bold=True, color="FFFFFF")
    ws["A1"].fill = header_fill("375623")
    ws["A1"].alignment = Alignment(horizontal="center")
    ws.merge_cells("A1:F1")

    ws["A2"] = "점검일: 2026-05-09 (목)"
    ws["D2"] = "라인: CVD-2라인"
    ws["A3"] = "점검자: 박설비"
    ws["D3"] = "승인: 최팀장"

    headers = ["No.", "설비명", "점검항목", "기준", "결과(O/X)", "비고"]
    for c, h in enumerate(headers, 1):
        cell = ws.cell(row=5, column=c, value=h)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = header_fill("548235")
        cell.alignment = Alignment(horizontal="center")
        cell.border = border_all

    items = [
        (1,  "CVD챔버 A",    "챔버압력 정상 여부",      "≤ 5 mTorr",  "O", ""),
        (2,  "CVD챔버 A",    "히터온도 편차 확인",       "±5°C 이내",  "O", ""),
        (3,  "CVD챔버 A",    "가스라인 누설 점검",       "누설 없음",   "X", "소량 누설 감지"),
        (4,  "로봇암 B",     "그리퍼 토크 점검",         "5 N·m 이상",  "O", ""),
        (5,  "로봇암 B",     "위치 재현성 확인",          "±0.1mm",    "O", ""),
        (6,  "컨베이어 C",   "벨트 장력 점검",            "정상 범위",   "O", ""),
        (7,  "컨베이어 C",   "속도 편차 확인",             "±2% 이내",  "O", ""),
        (8,  "진공펌프 D",   "오일 레벨 확인",             "MIN~MAX",   "O", ""),
        (9,  "진공펌프 D",   "진공도 달성 여부",           "≤ 1×10⁻³", "O", ""),
        (10, "냉각수시스템",  "유량 정상 여부",             "15 L/min",  "O", ""),
    ]

    for r, row_data in enumerate(items, 6):
        for c, val in enumerate(row_data, 1):
            cell = ws.cell(row=r, column=c, value=val)
            cell.alignment = Alignment(horizontal="center" if c in [1, 5] else "left")
            if c == 5 and val == "X":
                cell.font = Font(color="FF0000", bold=True)
            cell.border = border_all

    apply_border(ws, 5, 5 + len(items), 1, 6)
    set_col_widths(ws, [6, 14, 22, 14, 12, 18])
    wb.save(SAVE_DIR / "설비점검일지.xlsx")
    print("  [Excel 2/5] 설비점검일지.xlsx 생성 완료")

    # ── 3) 부품재고현황 ──────────────────────
    wb = Workbook()
    ws = wb.active
    ws.title = "부품재고현황"

    ws["A1"] = "부품 재고 현황"
    ws["A1"].font = Font(size=15, bold=True, color="FFFFFF")
    ws["A1"].fill = header_fill("7030A0")
    ws["A1"].alignment = Alignment(horizontal="center")
    ws.merge_cells("A1:H1")

    ws["A2"] = "기준일: 2026-05-09"
    ws["E2"] = "창고: 제1자재창고"

    headers = ["No.", "부품코드", "부품명", "규격", "단위", "현재고", "안전재고", "부족수량"]
    for c, h in enumerate(headers, 1):
        cell = ws.cell(row=4, column=c, value=h)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = header_fill("9E5DB7")
        cell.alignment = Alignment(horizontal="center")
        cell.border = border_all

    stock_data = [
        (1,  "SP-001", "볼베어링 6204",      "Φ20×Φ47×14",  "EA",  12,  10,  None),
        (2,  "SP-002", "타이밍벨트 T5-600",  "L=600mm",      "EA",   3,   5,  None),
        (3,  "SP-003", "실린더 패킹 세트",    "Φ50",          "SET",  8,   5,  None),
        (4,  "SP-004", "솔레노이드밸브 5/2방", "DC24V",        "EA",   1,   3,  None),
        (5,  "SP-005", "에어필터 엘리먼트",   "AFF40",        "EA",   0,   6,  None),
        (6,  "SP-006", "서보드라이버 파라미터 저장카드", "CF-4GB", "EA", 2,  2,  None),
        (7,  "SP-007", "히팅로드 Φ8×200",    "1kW",          "EA",   4,   4,  None),
        (8,  "SP-008", "써모커플 K타입",      "L=300mm",      "EA",   7,   5,  None),
        (9,  "SP-009", "냉각팬 모터",         "AC220V 25W",   "EA",   0,   2,  None),
        (10, "SP-010", "유압씰 세트",         "로드 Φ32",     "SET",  5,   4,  None),
    ]

    for r, row_data in enumerate(stock_data, 5):
        no, code, name, spec, unit, current, safety, _ = row_data
        shortage = max(0, safety - current) if current < safety else 0
        row_vals = [no, code, name, spec, unit, current, safety, shortage if shortage > 0 else ""]
        for c, val in enumerate(row_vals, 1):
            cell = ws.cell(row=r, column=c, value=val)
            cell.alignment = Alignment(horizontal="center" if c in [1, 5, 6, 7, 8] else "left")
            cell.border = border_all
            if c == 8 and val and val > 0:
                cell.fill = PatternFill("solid", fgColor="FFCCCC")
                cell.font = Font(color="CC0000", bold=True)
            if c == 6 and val == 0:
                cell.fill = PatternFill("solid", fgColor="FFCCCC")

    apply_border(ws, 4, 4 + len(stock_data), 1, 8)
    set_col_widths(ws, [6, 12, 22, 18, 6, 10, 10, 10])
    wb.save(SAVE_DIR / "부품재고현황.xlsx")
    print("  [Excel 3/5] 부품재고현황.xlsx 생성 완료")

    # ── 4) 월간정비실적 ──────────────────────
    wb = Workbook()
    ws = wb.active
    ws.title = "월간정비실적"

    ws["A1"] = "2026년 5월 월간 정비 실적"
    ws["A1"].font = Font(size=15, bold=True, color="FFFFFF")
    ws["A1"].fill = header_fill("C55A11")
    ws["A1"].alignment = Alignment(horizontal="center")
    ws.merge_cells("A1:I1")

    ws["A2"] = "보고기간: 2026-05-01 ~ 2026-05-09"
    ws["F2"] = "작성: 설비관리팀"

    headers = ["부서", "담당설비", "총정비건", "예방정비", "사후정비", "계획정비", "긴급정비", "가동률(%)", "비고"]
    for c, h in enumerate(headers, 1):
        cell = ws.cell(row=4, column=c, value=h)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = header_fill("ED7D31")
        cell.alignment = Alignment(horizontal="center")
        cell.border = border_all

    dept_data = [
        ("CVD공정팀",  "CVD챔버 A/B",    12, 8, 4, 7, 5, 96.2, ""),
        ("에칭공정팀", "에칭설비 1~3호",  9, 6, 3, 5, 4, 97.1, ""),
        ("세정팀",     "세정기 A/B/C",    7, 5, 2, 4, 3, 98.5, ""),
        ("물류팀",     "AGV, 컨베이어",  15, 9, 6, 8, 7, 94.8, "AGV #3 고장"),
        ("유틸리티팀", "냉각수, 진공계",   6, 4, 2, 3, 3, 99.1, ""),
        ("계측팀",     "검사장비 전체",    5, 3, 2, 3, 2, 98.0, ""),
    ]

    for r, row_data in enumerate(dept_data, 5):
        for c, val in enumerate(row_data, 1):
            cell = ws.cell(row=r, column=c, value=val)
            cell.alignment = Alignment(horizontal="center" if c not in [1, 2, 9] else "left")
            cell.border = border_all
            if c == 8:
                if val < 95:
                    cell.fill = PatternFill("solid", fgColor="FFCCCC")
                elif val >= 98:
                    cell.fill = PatternFill("solid", fgColor="CCFFCC")

    # 합계행
    total_row = 5 + len(dept_data)
    ws.cell(row=total_row, column=1, value="합계").font = Font(bold=True)
    for c in range(3, 8):
        col_vals = [ws.cell(row=r, column=c).value for r in range(5, total_row)]
        ws.cell(row=total_row, column=c, value=sum(col_vals)).font = Font(bold=True)
    avg_rate = sum(r[7] for r in dept_data) / len(dept_data)
    ws.cell(row=total_row, column=8, value=round(avg_rate, 1)).font = Font(bold=True)
    ws.cell(row=total_row, column=8).fill = PatternFill("solid", fgColor="FFF2CC")
    apply_border(ws, 4, total_row, 1, 9)
    set_col_widths(ws, [14, 18, 10, 10, 10, 10, 10, 12, 16])
    wb.save(SAVE_DIR / "월간정비실적.xlsx")
    print("  [Excel 4/5] 월간정비실적.xlsx 생성 완료")

    # ── 5) 작업지시서 ──────────────────────
    wb = Workbook()
    ws = wb.active
    ws.title = "작업지시서"

    ws["A1"] = "설비 정비 작업지시서"
    ws["A1"].font = Font(size=16, bold=True, color="FFFFFF")
    ws["A1"].fill = header_fill("1F3864")
    ws["A1"].alignment = Alignment(horizontal="center")
    ws.merge_cells("A1:F1")

    # 단건 양식 형태
    form_fields = [
        ("작업지시번호", "WO-2026-050901"),
        ("요청일",       "2026-05-09"),
        ("요청부서",     "CVD공정팀"),
        ("설비명",       "CVD챔버 A (EQ-A101)"),
        ("설비위치",     "C동 2층 CVD구역 A-03"),
        ("작업유형",     "사후정비 (Corrective Maintenance)"),
        ("긴급도",       "긴급 (4시간 이내 조치 필요)"),
        ("작업내용",     "챔버 상부 히터 교체 및 온도 캘리브레이션"),
        ("작업자",       "김정비 (주), 이보조 (보조)"),
        ("완료예정일",   "2026-05-09 18:00"),
        ("특이사항",     "작업 전 챔버 내부 잔류가스 퍼지 필수"),
        ("결재선",       "작성: 박대리 → 검토: 최과장 → 승인: 이부장"),
    ]

    for r, (label, value) in enumerate(form_fields, 3):
        label_cell = ws.cell(row=r, column=1, value=label)
        label_cell.font = Font(bold=True)
        label_cell.fill = PatternFill("solid", fgColor="D9E1F2")
        label_cell.alignment = Alignment(horizontal="right")
        label_cell.border = border_all

        colon_cell = ws.cell(row=r, column=2, value=":")
        colon_cell.alignment = Alignment(horizontal="center")
        colon_cell.border = border_all

        val_cell = ws.cell(row=r, column=3, value=value)
        val_cell.alignment = Alignment(horizontal="left")
        val_cell.border = border_all
        ws.merge_cells(start_row=r, start_column=3, end_row=r, end_column=6)

    set_col_widths(ws, [18, 4, 20, 15, 15, 15])
    wb.save(SAVE_DIR / "작업지시서.xlsx")
    print("  [Excel 5/5] 작업지시서.xlsx 생성 완료")


# ──────────────────────────────────────────────
# 2. Word 샘플 5개
# ──────────────────────────────────────────────
def create_word_samples():
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def set_heading(doc, text, level=1, color=None):
        p = doc.add_heading(text, level=level)
        if color:
            for run in p.runs:
                run.font.color.rgb = RGBColor(*color)
        return p

    def add_info_table(doc, rows):
        """2열 정보 테이블 (항목 | 내용)"""
        table = doc.add_table(rows=len(rows), cols=2)
        table.style = "Table Grid"
        for i, (label, value) in enumerate(rows):
            label_cell = table.cell(i, 0)
            label_cell.text = label
            label_cell.paragraphs[0].runs[0].font.bold = True
            table.cell(i, 1).text = value
        return table

    # ── 1) 설비일보_2026-05 ──────────────────────
    doc = Document()
    doc.add_heading("설비 일보", 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.add_run("작성일: 2026년 5월 9일 (목)    작성자: 박설비    라인: CVD-2라인").bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("1. 일반 현황", 1)
    add_info_table(doc, [
        ("라인명",     "CVD-2라인"),
        ("총 설비 수", "12대"),
        ("정상 가동",  "10대"),
        ("고장 중",    "1대 (CVD챔버 A)"),
        ("PM 중",      "1대 (로봇암 B)"),
        ("가동률",     "91.7%"),
    ])

    doc.add_heading("2. 고장 내역", 1)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["설비명", "발생시각", "고장내용", "가동시간(h)", "담당자"]
    for i, h in enumerate(headers):
        table.cell(0, i).text = h
        table.cell(0, i).paragraphs[0].runs[0].font.bold = True

    fault_data = [
        ("CVD챔버 A", "09:23", "히터 온도 편차 과대 (±15°C)", "6.4", "김정비"),
        ("AGV #7",    "13:47", "배터리 저전압 경보",           "4.8", "이보조"),
    ]
    for row_data in fault_data:
        row = table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = val

    doc.add_heading("3. 고장 분석", 1)
    for title, content in [
        ("현상", "CVD챔버 A 히터 영역별 온도 편차가 설정값 대비 ±15°C 초과 발생. 정상 기준 ±5°C 이내."),
        ("원인", "히터 로드 #3번 열화로 인한 발열 불균일. 최근 PM 주기(3개월) 경과로 교체 시기 도래."),
        ("조치", "히터 로드 #3번 교체 완료 (SP-007 사용). 온도 캘리브레이션 후 정상 확인. 잔여 히터 로드 예방 점검 일정 수립."),
    ]:
        p = doc.add_paragraph()
        p.add_run(f"[{title}] ").bold = True
        p.add_run(content)

    doc.add_heading("4. 익일 작업 계획", 1)
    items = [
        "로봇암 B PM 완료 및 정상 가동 복귀 (오전)",
        "CVD챔버 B 예방점검 (오후 2시 예정)",
        "진공펌프 오일 교환 (PM 주기 도래)",
    ]
    for item in items:
        doc.add_paragraph(item, style="List Bullet")

    doc.save(SAVE_DIR / "설비일보_2026-05.docx")
    print("  [Word 1/5] 설비일보_2026-05.docx 생성 완료")

    # ── 2) 정비작업보고서 ──────────────────────
    doc = Document()
    doc.add_heading("정비 작업 보고서", 0).alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_info_table(doc, [
        ("보고서 번호", "REP-2026-050901"),
        ("작업일시",    "2026년 5월 9일 09:30 ~ 15:00"),
        ("작업자",      "김정비 (주임), 이보조 (사원)"),
        ("설비명",      "CVD챔버 A (EQ-A101)"),
        ("설비위치",    "C동 2층 CVD구역 A-03"),
        ("작업유형",    "사후정비"),
    ])

    doc.add_heading("작업 내용", 1)
    doc.add_paragraph(
        "CVD챔버 A의 히터 온도 편차 과대 이상으로 인해 히터 로드 #3번 교체 작업을 수행하였다. "
        "작업 착수 전 챔버 내부 잔류 가스 퍼지를 완료하고, 안전 잠금(LOTO)을 적용하였다. "
        "히터 로드 탈거 후 육안 점검에서 열화에 의한 발열체 단선이 확인되어 신품으로 교체하였다. "
        "교체 완료 후 온도 캘리브레이션을 수행하여 전 영역 ±3°C 이내 달성을 확인하였다."
    )

    doc.add_heading("사용 부품", 1)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    for i, h in enumerate(["No.", "부품코드", "부품명", "수량", "비고"]):
        table.cell(0, i).text = h
        table.cell(0, i).paragraphs[0].runs[0].font.bold = True

    parts = [
        ("1", "SP-007", "히팅로드 Φ8×200 (1kW)", "1", "재고 사용"),
        ("2", "SP-003", "실린더 패킹 세트 Φ50",   "1", "재고 사용"),
        ("3", "-",      "내열 실리콘 그리스",       "소량", "현장 보유분"),
    ]
    for p_data in parts:
        row = table.add_row()
        for i, val in enumerate(p_data):
            row.cells[i].text = val

    doc.add_heading("작업 결과", 1)
    for label, text in [
        ("결과",     "정상 복구 완료"),
        ("확인사항", "히터 전영역 온도 균일성 ±3°C 이내 달성 확인"),
        ("재발방지", "히터 로드 PM 주기를 3개월에서 2개월로 단축. 예비 히터 로드 재고 보충 요청."),
    ]:
        p = doc.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run(text)

    doc.save(SAVE_DIR / "정비작업보고서.docx")
    print("  [Word 2/5] 정비작업보고서.docx 생성 완료")

    # ── 3) 설비이력카드 ──────────────────────
    doc = Document()
    doc.add_heading("설비 이력 카드", 0).alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_info_table(doc, [
        ("설비번호",   "EQ-A101"),
        ("설비명",     "CVD 챔버 A (Chemical Vapor Deposition)"),
        ("제조사",     "(주)한국설비기계"),
        ("모델명",     "CVD-5000S"),
        ("설치일",     "2022년 3월 15일"),
        ("담당부서",   "CVD공정팀"),
        ("담당자",     "김정비 주임"),
        ("설치위치",   "C동 2층 CVD구역 A-03"),
        ("사용유틸리티", "전기 (AC 380V, 60kW), 냉각수 (15 L/min), 질소 (2 kg/cm²)"),
    ])

    doc.add_heading("주요 정비 이력", 1)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    for i, h in enumerate(["일자", "작업유형", "작업내용", "사용부품", "담당자"]):
        table.cell(0, i).text = h
        table.cell(0, i).paragraphs[0].runs[0].font.bold = True

    history = [
        ("2022-03-15", "설치",   "초기 설치 및 시운전",                     "-",          "이과장"),
        ("2022-09-01", "예방정비", "전체 PM (히터, 씰 교체)",               "SP-003,007", "김정비"),
        ("2023-03-10", "예방정비", "반기 PM",                                "SP-003",     "김정비"),
        ("2023-08-22", "사후정비", "진공펌프 오일누설 수리",                 "SP-010",     "박대리"),
        ("2024-01-05", "예방정비", "연간 PM (히터, 씰, 베어링 전체 교체)",  "SP-001,003,007", "김정비"),
        ("2024-07-17", "사후정비", "솔레노이드밸브 고착 교체",               "SP-004",     "이보조"),
        ("2025-02-28", "예방정비", "반기 PM",                                "SP-003,007", "김정비"),
        ("2025-09-11", "사후정비", "온도센서 단선 교체",                     "SP-008",     "김정비"),
        ("2026-05-09", "사후정비", "히터 로드 #3번 교체, 온도 캘리브레이션", "SP-007",     "김정비"),
    ]
    for h_data in history:
        row = table.add_row()
        for i, val in enumerate(h_data):
            row.cells[i].text = val

    doc.save(SAVE_DIR / "설비이력카드.docx")
    print("  [Word 3/5] 설비이력카드.docx 생성 완료")

    # ── 4) 안전점검결과서 ──────────────────────
    doc = Document()
    doc.add_heading("안전 점검 결과서", 0).alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_info_table(doc, [
        ("점검일",   "2026년 5월 9일"),
        ("점검자",   "안전관리팀 최안전 대리"),
        ("점검구역", "C동 2층 CVD공정 구역"),
        ("동행자",   "CVD공정팀 김정비 주임"),
    ])

    doc.add_heading("점검 항목별 결과", 1)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    for i, h in enumerate(["No.", "점검항목", "점검기준", "결과", "개선조치"]):
        table.cell(0, i).text = h
        table.cell(0, i).paragraphs[0].runs[0].font.bold = True

    check_items = [
        ("1",  "소화기 배치",       "각 구역 5m 이내 배치",    "적합",   "-"),
        ("2",  "비상구 표시",       "형광 표시 선명",           "적합",   "-"),
        ("3",  "전기 배전반 상태",  "문 잠금, 경고표지 부착",   "적합",   "-"),
        ("4",  "가스배관 표지",     "색상 코드 및 방향 표시",   "부적합", "5/14까지 표지 보완 완료 예정"),
        ("5",  "안전화 착용",       "작업자 전원 착용",         "적합",   "-"),
        ("6",  "LOTO 키 관리",      "지정함 보관, 열쇠 이중화", "적합",   "-"),
        ("7",  "케미컬 보관",       "밀폐용기, 라벨 부착",      "적합",   "-"),
        ("8",  "환기 시스템",       "가동 상태, 필터 청결",     "부적합", "필터 교체 즉시 조치 완료"),
        ("9",  "비상정지 버튼",     "각 설비 1개 이상 설치",   "적합",   "-"),
        ("10", "바닥 안전선 도색",  "황색선 선명, 박리 없음",   "적합",   "-"),
    ]
    for c_data in check_items:
        row = table.add_row()
        for i, val in enumerate(c_data):
            row.cells[i].text = val

    doc.add_heading("종합 의견", 1)
    doc.add_paragraph(
        "금번 안전점검 결과 총 10개 항목 중 2개 항목에서 부적합 사항이 확인되었습니다. "
        "환기 시스템 필터 교체는 점검 당일 즉시 조치 완료하였으며, "
        "가스배관 표지 보완은 2026년 5월 14일까지 완료 예정으로 담당자 확인을 받았습니다. "
        "2주 후 재점검을 통해 개선 완료 여부를 확인할 예정입니다."
    )

    doc.save(SAVE_DIR / "안전점검결과서.docx")
    print("  [Word 4/5] 안전점검결과서.docx 생성 완료")

    # ── 5) 교육훈련기록 ──────────────────────
    doc = Document()
    doc.add_heading("교육 훈련 기록부", 0).alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_info_table(doc, [
        ("교육명",   "CVD 설비 긴급조치 절차 교육"),
        ("일시",     "2026년 5월 8일 14:00 ~ 16:00 (2시간)"),
        ("장소",     "C동 3층 교육실"),
        ("강사",     "김정비 주임 (CVD공정팀)"),
        ("총 참석자", "12명"),
    ])

    doc.add_heading("교육 내용", 1)
    contents = [
        "1. CVD 공정 개요 및 주요 설비 구성 (30분)",
        "2. 이상 유형별 1차 대응 절차 (40분)",
        "   - 온도 이상, 진공 이상, 가스 누설 시 즉각 대응",
        "   - 비상정지 순서 및 LOTO 적용 방법",
        "3. 실습: 히터 이상 시뮬레이션 대응 (30분)",
        "4. Q&A 및 평가 (20분)",
    ]
    for c in contents:
        doc.add_paragraph(c)

    doc.add_heading("참석자 명단", 1)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    for i, h in enumerate(["No.", "성명", "부서", "직급", "서명"]):
        table.cell(0, i).text = h
        table.cell(0, i).paragraphs[0].runs[0].font.bold = True

    attendees = [
        ("1",  "박설비", "CVD공정팀",  "사원",  ""),
        ("2",  "이보조", "CVD공정팀",  "사원",  ""),
        ("3",  "최운영", "CVD공정팀",  "대리",  ""),
        ("4",  "정안전", "안전관리팀", "대리",  ""),
        ("5",  "한점검", "설비관리팀", "주임",  ""),
        ("6",  "강신입", "CVD공정팀",  "사원",  ""),
        ("7",  "윤경험", "에칭공정팀", "대리",  ""),
        ("8",  "임교대", "CVD공정팀",  "사원",  ""),
        ("9",  "신협력", "협력업체",   "기사",  ""),
        ("10", "조지원", "CVD공정팀",  "사원",  ""),
        ("11", "권현장", "물류팀",     "주임",  ""),
        ("12", "문배우", "신입사원",   "수습",  ""),
    ]
    for a_data in attendees:
        row = table.add_row()
        for i, val in enumerate(a_data):
            row.cells[i].text = val

    doc.add_heading("교육 평가 결과", 1)
    add_info_table(doc, [
        ("평가 방법", "필기시험 (10문항)"),
        ("평균 점수", "87.5점"),
        ("최고 점수", "100점 (박설비)"),
        ("미달자 (70점 미만)", "없음"),
        ("재교육 필요자", "없음"),
    ])

    doc.save(SAVE_DIR / "교육훈련기록.docx")
    print("  [Word 5/5] 교육훈련기록.docx 생성 완료")


# ──────────────────────────────────────────────
# 3. PDF 샘플 5개
# ──────────────────────────────────────────────
def create_pdf_samples():
    from fpdf import FPDF

    # fpdf2에서 한국어를 쓰려면 유니코드 폰트 필요.
    # 시스템 폰트(malgun.ttf) 또는 내장 latin-1 폰트로 fallback.
    import shutil, sys

    FONT_PATHS = [
        "C:/Windows/Fonts/malgun.ttf",
        "C:/Windows/Fonts/NanumGothic.ttf",
        "C:/Windows/Fonts/gulim.ttc",
    ]
    FONT_FILE = None
    for fp in FONT_PATHS:
        if os.path.exists(fp):
            FONT_FILE = fp
            break

    class KorPDF(FPDF):
        """한국어 지원 FPDF 래퍼"""
        def __init__(self):
            super().__init__()
            self.kor_font = "kor"
            self._kor_available = False
            if FONT_FILE:
                try:
                    # fpdf2 v2.5.1+ : uni 파라미터 불필요, bold 별도 등록
                    self.add_font(self.kor_font, style="", fname=FONT_FILE)
                    # Bold는 같은 폰트로 등록 (굴림/맑은고딕은 bold 전용 파일 없음)
                    self.add_font(self.kor_font, style="B", fname=FONT_FILE)
                    self._kor_available = True
                except Exception:
                    self.kor_font = "Helvetica"

        def set_kor(self, size=10, bold=False):
            style = "B" if bold else ""
            if self._kor_available:
                self.set_font(self.kor_font, style, size)
            else:
                self.set_font("Helvetica", style, size)

        def title_box(self, text, size=16):
            self.set_fill_color(31, 78, 121)
            self.set_text_color(255, 255, 255)
            self.set_kor(size, bold=True)
            self.cell(0, 12, text, new_x="LMARGIN", new_y="NEXT", align="C", fill=True)
            self.set_text_color(0, 0, 0)
            self.ln(4)

        def section(self, text):
            self.set_fill_color(214, 220, 229)
            self.set_kor(11, bold=True)
            self.cell(0, 8, text, new_x="LMARGIN", new_y="NEXT", fill=True)
            self.ln(1)

        def kv(self, label, value, label_w=50):
            self.set_kor(9, bold=True)
            self.cell(label_w, 6, label)
            self.set_kor(9)
            self.cell(0, 6, value, new_x="LMARGIN", new_y="NEXT")

        def table_header(self, cols, widths):
            self.set_fill_color(46, 117, 182)
            self.set_text_color(255, 255, 255)
            self.set_kor(9, bold=True)
            for col, w in zip(cols, widths):
                self.cell(w, 7, col, border=1, align="C", fill=True)
            self.ln()
            self.set_text_color(0, 0, 0)

        def table_row(self, vals, widths, aligns=None, fill=False):
            if fill:
                self.set_fill_color(242, 242, 242)
            self.set_kor(9)
            if aligns is None:
                aligns = ["L"] * len(vals)
            for val, w, align in zip(vals, widths, aligns):
                self.cell(w, 6, str(val), border=1, align=align, fill=fill)
            self.ln()

    # ── 1) 품질검사성적서 ──────────────────────
    pdf = KorPDF()
    pdf.add_page()
    pdf.title_box("품질 검사 성적서")

    pdf.kv("성적서 번호:", "QC-2026-050901")
    pdf.kv("발행일:",      "2026-05-09")
    pdf.kv("제품명:",      "OLED 패널 A-Type (65인치)")
    pdf.kv("LOT 번호:",    "LOT-2026-05-001")
    pdf.kv("생산라인:",    "CVD-2라인")
    pdf.kv("검사자:",      "품질팀 홍검사")
    pdf.kv("판정:",        "합격 (PASS)")
    pdf.ln(4)

    pdf.section("검사 항목별 결과")
    cols   = ["No.", "검사항목",          "단위",   "규격(하한)", "규격(상한)", "측정치",  "판정"]
    widths = [10,    55,                   15,       22,            22,            25,         20]
    aligns = ["C",   "L",                  "C",      "C",           "C",           "C",        "C"]
    pdf.table_header(cols, widths)

    items = [
        ("1", "박막두께",           "nm",    "95.0",  "105.0", "101.2", "합격"),
        ("2", "면저항",             "Ω/sq",  "8.0",   "12.0",  "9.8",   "합격"),
        ("3", "표면 조도 (Ra)",     "nm",    "0",     "2.0",   "1.3",   "합격"),
        ("4", "접착력",             "N/cm",  "5.0",   "-",     "7.2",   "합격"),
        ("5", "투과율",             "%",     "88.0",  "92.0",  "90.1",  "합격"),
        ("6", "색도 (CIE x)",       "-",     "0.310", "0.330", "0.319", "합격"),
        ("7", "색도 (CIE y)",       "-",     "0.320", "0.340", "0.328", "합격"),
        ("8", "파티클 수 (≥0.5μm)", "ea/m²", "0",     "50",    "12",    "합격"),
    ]
    for idx, row in enumerate(items):
        pdf.table_row(row, widths, aligns, fill=(idx % 2 == 1))

    pdf.output(str(SAVE_DIR / "품질검사성적서.pdf"))
    print("  [PDF 1/5] 품질검사성적서.pdf 생성 완료")

    # ── 2) 외주정비계약서 ──────────────────────
    pdf = KorPDF()
    pdf.add_page()
    pdf.title_box("외주 정비 계약서")

    pdf.kv("계약번호:",   "CNT-2026-0509-001")
    pdf.kv("계약일:",     "2026년 5월 9일")
    pdf.kv("발주처:",     "(주)디스플레이코리아 설비관리팀")
    pdf.kv("수주처:",     "(주)정비전문서비스")
    pdf.kv("계약기간:",   "2026년 6월 1일 ~ 2027년 5월 31일 (1년)")
    pdf.kv("계약금액:",   "￦ 48,000,000 (부가세 별도)")
    pdf.ln(4)

    pdf.section("계약 대상 설비")
    cols   = ["No.", "설비명",             "수량", "정비주기", "단가(원/회)"]
    widths = [10,    70,                    15,     30,          45]
    aligns = ["C",   "L",                   "C",    "C",         "R"]
    pdf.table_header(cols, widths)
    equip = [
        ("1", "CVD챔버 A/B",       "2대", "월 1회",  "1,200,000"),
        ("2", "에칭설비 1~3호",    "3대", "월 1회",    "900,000"),
        ("3", "진공펌프 A/B/C/D",  "4대", "분기 1회",  "600,000"),
        ("4", "로봇암 B, C",       "2대", "반기 1회",  "800,000"),
    ]
    for idx, row in enumerate(equip):
        pdf.table_row(row, widths, aligns, fill=(idx % 2 == 1))

    pdf.ln(4)
    pdf.section("특약 사항")
    specials = [
        "1. 계약 기간 중 긴급출동 (4시간 이내) 서비스를 월 2회 무상 제공한다.",
        "2. 수주처는 정비 완료 후 48시간 이내 결과보고서를 제출하여야 한다.",
        "3. 정비 불량으로 인한 재발 고장은 수주처 비용으로 재정비한다.",
        "4. 정비 담당자는 자격증 보유자로 한정하며, 변경 시 사전 승인을 받아야 한다.",
        "5. 계약 해지 시 1개월 전 서면 통보를 원칙으로 한다.",
    ]
    pdf.set_kor(9)
    for s in specials:
        pdf.set_x(pdf.l_margin)
        pdf.multi_cell(0, 6, s)

    pdf.output(str(SAVE_DIR / "외주정비계약서.pdf"))
    print("  [PDF 2/5] 외주정비계약서.pdf 생성 완료")

    # ── 3) 설비도면목록 ──────────────────────
    pdf = KorPDF()
    pdf.add_page()
    pdf.title_box("설비 도면 목록")

    pdf.kv("관리부서:", "설비관리팀")
    pdf.kv("작성일:",   "2026-05-09")
    pdf.kv("문서번호:", "DWG-LIST-2026-001")
    pdf.kv("개정:",     "Rev.3")
    pdf.ln(4)

    pdf.section("CVD설비 도면 목록")
    cols   = ["No.", "도면번호",          "도면명",                "버전", "등록일",     "담당자", "비고"]
    widths = [8,     32,                   55,                       12,     22,            18,       23]
    aligns = ["C",   "L",                  "L",                      "C",    "C",           "C",      "L"]
    pdf.table_header(cols, widths)

    drawings = [
        ("1",  "DWG-CVD-001", "CVD챔버 A 전체 조립도",          "V3.1", "2022-03-10", "김설계", "현행"),
        ("2",  "DWG-CVD-002", "CVD챔버 A 히터 상세도",          "V2.0", "2023-01-05", "김설계", "현행"),
        ("3",  "DWG-CVD-003", "CVD챔버 A 진공계통 P&ID",        "V1.5", "2022-06-20", "이배관", "현행"),
        ("4",  "DWG-CVD-004", "CVD챔버 B 전체 조립도",          "V2.3", "2022-03-10", "김설계", "현행"),
        ("5",  "DWG-CVD-005", "가스공급 배관 계통도",            "V4.0", "2025-02-15", "이배관", "현행"),
        ("6",  "DWG-ROB-001", "로봇암 B 기구도",                "V1.2", "2022-04-01", "박기구", "현행"),
        ("7",  "DWG-ROB-002", "로봇암 C 기구도",                "V1.1", "2022-04-01", "박기구", "현행"),
        ("8",  "DWG-VAC-001", "진공펌프 계통도",                "V2.0", "2023-09-10", "최설계", "현행"),
        ("9",  "DWG-CVY-001", "컨베이어 C 레이아웃",            "V1.0", "2022-05-20", "박기구", "현행"),
        ("10", "DWG-UTL-001", "냉각수 계통 전체도",             "V3.2", "2025-07-30", "이배관", "현행"),
    ]
    for idx, row in enumerate(drawings):
        pdf.table_row(row, widths, aligns, fill=(idx % 2 == 1))

    pdf.output(str(SAVE_DIR / "설비도면목록.pdf"))
    print("  [PDF 3/5] 설비도면목록.pdf 생성 완료")

    # ── 4) 사고보고서 ──────────────────────
    pdf = KorPDF()
    pdf.add_page()
    pdf.title_box("사고 보고서")

    pdf.kv("보고서 번호:", "ACC-2026-050901")
    pdf.kv("발생일시:",    "2026년 5월 9일 09:23")
    pdf.kv("보고일시:",    "2026년 5월 9일 10:05")
    pdf.kv("보고자:",      "CVD공정팀 박설비 사원")
    pdf.kv("승인자:",      "CVD공정팀 이과장")
    pdf.ln(4)

    for title, content in [
        ("1. 사고 개요",
         "2026년 5월 9일 09:23경 C동 2층 CVD구역 A-03에 위치한 CVD챔버 A에서\n"
         "히터 온도 편차 과대 경보가 발생하였다. 공정 중단 및 안전 조치를 취하였으며,\n"
         "인명 피해 및 제품 피해는 없었다."),
        ("2. 사고 발생 장소 및 관련 설비",
         "장소: C동 2층 CVD구역 A-03\n"
         "관련 설비: CVD챔버 A (EQ-A101, 설치일 2022-03-15)\n"
         "관련 공정: OLED 패널 CVD 박막 증착 공정"),
        ("3. 사고 내용 (현상)",
         "CVD 공정 진행 중 히터 영역 #3의 온도가 설정값 대비 +18°C 초과 상승.\n"
         "설비 자동 인터락 작동으로 공정 중단. 이후 챔버 내부 온도 불균일로\n"
         "생산 중이던 LOT-2026-05-001 패널 4매 재작업 결정."),
        ("4. 원인 분석",
         "히터 로드 #3번의 열화로 인한 발열 불균일.\n"
         "해당 히터 로드는 설치 이후 약 36개월 사용으로 PM 교체 주기 경과.\n"
         "최근 PM 미실시 (예산 이슈로 2개월 지연)."),
        ("5. 조치 사항",
         "- 즉각 조치: 공정 중단, LOTO 적용, 안전 확인\n"
         "- 당일 조치: 히터 로드 #3번 교체 (SP-007) 및 온도 캘리브레이션\n"
         "- 교체 후 전 영역 온도 균일성 ±3°C 이내 확인, 오후 15:30 정상 가동 재개"),
        ("6. 재발 방지 대책",
         "- 히터 로드 PM 주기: 3개월 → 2개월으로 단축\n"
         "- 히터 예비 부품 재고 상시 보유 (최소 2ea) 정책 수립\n"
         "- PM 지연 발생 시 즉시 팀장 보고 프로세스 강화\n"
         "- 온도 편차 경보 기준 강화: ±15°C → ±10°C"),
    ]:
        pdf.section(title)
        pdf.set_kor(9)
        pdf.multi_cell(0, 5.5, content)
        pdf.ln(2)

    pdf.output(str(SAVE_DIR / "사고보고서.pdf"))
    print("  [PDF 4/5] 사고보고서.pdf 생성 완료")

    # ── 5) 자재발주요청서 ──────────────────────
    pdf = KorPDF()
    pdf.add_page()
    pdf.title_box("자재 발주 요청서")

    pdf.kv("요청번호:", "PO-REQ-2026-050901")
    pdf.kv("요청일:",   "2026-05-09")
    pdf.kv("요청부서:", "설비관리팀")
    pdf.kv("요청자:",   "김정비 주임")
    pdf.kv("긴급도:",   "일반 (납기 5 영업일 이내)")
    pdf.kv("배송처:",   "C동 1층 자재창고 (창고 담당: 홍창고)")
    pdf.ln(4)

    pdf.section("자재 목록")
    cols   = ["No.", "부품코드", "부품명",              "규격",          "단위", "수량", "단가(원)",   "합계(원)"]
    widths = [8,     20,         48,                     35,              10,     10,     22,            22]
    aligns = ["C",   "L",        "L",                    "L",             "C",    "C",    "R",           "R"]
    pdf.table_header(cols, widths)

    items = [
        ("1",  "SP-007", "히팅로드",             "Φ8×200 (1kW)",     "EA",  3,  35000,   105000),
        ("2",  "SP-003", "실린더 패킹 세트",      "Φ50",              "SET", 5,  18000,    90000),
        ("3",  "SP-004", "솔레노이드밸브 5/2방",  "DC24V",            "EA",  2,  98000,   196000),
        ("4",  "SP-005", "에어필터 엘리먼트",     "AFF40",            "EA",  6,  25000,   150000),
        ("5",  "SP-009", "냉각팬 모터",           "AC220V 25W",       "EA",  2, 185000,   370000),
        ("6",  "SP-001", "볼베어링 6204",         "Φ20×Φ47×14",      "EA", 10,   8500,    85000),
        ("7",  "SP-008", "써모커플 K타입",        "L=300mm",          "EA",  4,  22000,    88000),
        ("8",  "SP-010", "유압씰 세트",           "로드 Φ32",         "SET", 3,  31000,    93000),
    ]
    total = 0
    for idx, row in enumerate(items):
        no, code, name, spec, unit, qty, unit_price, amount = row
        pdf.table_row(
            [no, code, name, spec, unit, str(qty),
             f"{unit_price:,}", f"{amount:,}"],
            widths, aligns, fill=(idx % 2 == 1)
        )
        total += amount

    # 합계행
    pdf.set_fill_color(255, 242, 204)
    pdf.set_kor(9, bold=True)
    merged_w = sum(widths[:-1])
    pdf.cell(merged_w, 7, "합계", border=1, align="R", fill=True)
    pdf.cell(widths[-1], 7, f"{total:,}", border=1, align="R", fill=True)
    pdf.ln()

    pdf.ln(4)
    pdf.section("비고")
    pdf.set_kor(9)
    pdf.multi_cell(0, 5.5,
        "- SP-005, SP-009는 재고 소진으로 긴급 발주 요망.\n"
        "- 납품 시 거래명세서 및 품질성적서(합격증) 동봉 필수.\n"
        "- 발주처 변경 시 구매팀 통해 재발주 요청할 것."
    )

    pdf.output(str(SAVE_DIR / "자재발주요청서.pdf"))
    print("  [PDF 5/5] 자재발주요청서.pdf 생성 완료")


# ──────────────────────────────────────────────
# 메인
# ──────────────────────────────────────────────
if __name__ == "__main__":
    print("=" * 60)
    print("  디스플레이 공장 정비/설비 관리 샘플 파일 생성")
    print(f"  저장 경로: {SAVE_DIR.resolve()}")
    print("=" * 60)

    print("\n[Excel 파일 생성 중...]")
    create_excel_samples()

    print("\n[Word 파일 생성 중...]")
    create_word_samples()

    print("\n[PDF 파일 생성 중...]")
    create_pdf_samples()

    print("\n" + "=" * 60)
    print("  생성 완료 - 파일 목록:")
    print("=" * 60)
    files = sorted(SAVE_DIR.glob("*"))
    total_size = 0
    counts = {"xlsx": 0, "docx": 0, "pdf": 0}
    for f in files:
        if f.is_file() and f.suffix in {".xlsx", ".docx", ".pdf"}:
            size = f.stat().st_size
            total_size += size
            counts[f.suffix.lstrip(".")] += 1
            print(f"  {f.name:<35}  {size:>8,} bytes")
    print("-" * 60)
    print(f"  Excel: {counts['xlsx']}개  |  Word: {counts['docx']}개  |  PDF: {counts['pdf']}개")
    print(f"  전체 {sum(counts.values())}개 파일  |  총 {total_size:,} bytes")
    print("=" * 60)
